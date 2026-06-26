#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
软著申请 —— 源代码 Word 文档生成器
====================================

根据中国版权保护中心官方规范，生成格式合规的源代码 Word 文档：
  - A4 纸，四边 2.5cm 页边距
  - 页眉：软件全称+版本号
  - 页脚：连续页码 "第X页"
  - 字体 Courier New，小五号 (9pt)
  - 单倍行距，段前段后0磅
  - 每页 ≥ 50 行代码
  - 前30页 + 后30页 = 60页
  - 首页为程序入口代码，末页为模块结束代码

输出：soft-copyright/output/软著源代码.docx
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ============ 配置 ============

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, '..'))

SOFTWARE_NAME = '码上成长Python智能学习平台'
VERSION = 'V1.0'
HEADER_TEXT = f'{SOFTWARE_NAME} {VERSION}'
LINES_PER_PAGE = 50
FRONT_PAGES = 30
BACK_PAGES = 30
CODE_FONT = 'Courier New'
CODE_FONT_SIZE = Pt(9)    # 小五号 ≈ 9pt
HEADER_FONT = '宋体'
HEADER_FONT_SIZE = Pt(10.5)  # 五号

# 源码目录与排除规则
INCLUDE_DIRS = ['src/app', 'src/lib', 'src/components', 'src/hooks']
EXCLUDE_DIRS = ['src/components/ui', 'node_modules', '.next', 'dist', 'data']
INCLUDE_EXTENSIONS = ['.ts', '.tsx']

# 文件优先级排序
DIR_PRIORITY = [
    'src/lib',
    'src/app/api',
    'src/app/login',
    'src/app/evaluate',
    'src/app/practice',
    'src/app/chat',
    'src/app/statistics',
    'src/app/resources',
    'src/app/admin',
    'src/app',
    'src/components',
    'src/hooks',
]

OUTPUT_DIR = os.path.join(SCRIPT_DIR, 'output')
OUTPUT_FILE = os.path.join(OUTPUT_DIR, '软著源代码.docx')


def should_exclude(path):
    normalized = path.replace('\\', '/')
    for excl in EXCLUDE_DIRS:
        if excl.replace('\\', '/') in normalized:
            return True
    return False


def collect_source_files():
    """按优先级收集源文件"""
    files_by_dir = {}
    for include_dir in INCLUDE_DIRS:
        full_dir = os.path.join(PROJECT_ROOT, include_dir)
        if not os.path.isdir(full_dir):
            continue
        for root, dirs, files in os.walk(full_dir):
            dirs[:] = [d for d in dirs if not should_exclude(os.path.join(root, d))]
            for fname in files:
                if not any(fname.endswith(ext) for ext in INCLUDE_EXTENSIONS):
                    continue
                full_path = os.path.join(root, fname)
                if should_exclude(full_path):
                    continue
                rel_path = os.path.relpath(full_path, PROJECT_ROOT)
                group = 'other'
                rel_slash = rel_path.replace('\\', '/')
                for pd in DIR_PRIORITY:
                    if rel_slash.startswith(pd.replace('\\', '/')):
                        group = pd
                        break
                if group not in files_by_dir:
                    files_by_dir[group] = []
                files_by_dir[group].append((rel_path, full_path))

    result = []
    for group in DIR_PRIORITY:
        if group in files_by_dir:
            files_by_dir[group].sort()
            result.extend(files_by_dir[group])
    if 'other' in files_by_dir:
        files_by_dir['other'].sort()
        result.extend(files_by_dir['other'])
    return result


def read_file_safe(filepath):
    encodings = ['utf-8', 'gbk', 'latin-1']
    for enc in encodings:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                return f.read().splitlines()
        except (UnicodeDecodeError, UnicodeError):
            continue
    try:
        with open(filepath, 'rb') as f:
            return f.read().decode('utf-8', errors='replace').splitlines()
    except Exception:
        return [f'# [读取失败: {os.path.basename(filepath)}]']


def clean_code_line(line):
    """清理代码行：删除开源协议声明、多余空行"""
    # 删除包含开源协议的行
    lower = line.lower()
    if any(kw in lower for kw in ['mit license', 'apache license', 'gpl', 'bsd license',
                                    'copyright (c)', 'all rights reserved',
                                    'this software is provided',
                                    'licensed under the mit',
                                    'licensed under the apache',
                                    'permission is hereby granted']):
        return ''  # 替换为空行
    # 去掉行尾空格
    return line.rstrip()


def build_continuous_code_blocks(files):
    """
    构建连续的代码块列表。
    每个文件之间用文件路径注释分隔，删除多余空行和开源声明。
    返回: all_lines (所有有效代码行的列表), file_markers (文件起始标记列表)
    """
    all_lines = []
    file_markers = []  # (line_index, rel_path)

    for rel_path, abs_path in files:
        lines = read_file_safe(abs_path)
        rel_slash = rel_path.replace('\\', '/')
        # 添加文件起始注释（3行，确保50行内不占太多）
        marker_idx = len(all_lines)
        all_lines.append(f'// ==============================')
        all_lines.append(f'// 文件: {rel_slash}')
        all_lines.append(f'// ==============================')
        file_markers.append((marker_idx, rel_slash))

        prev_empty = False
        for line in lines:
            cleaned = clean_code_line(line)
            # 去除连续空行（最多保留1个空行）
            if cleaned == '' or cleaned.strip() == '':
                if prev_empty:
                    continue  # 跳过连续空行
                prev_empty = True
                all_lines.append('')
            else:
                prev_empty = False
                all_lines.append(cleaned)

    return all_lines, file_markers


def paginate_lines(all_lines, lines_per_page):
    """将代码行按每页 lines_per_page 行分页"""
    pages = []
    idx = 0
    while idx < len(all_lines):
        page_lines = all_lines[idx:idx + lines_per_page]
        # 补空行到 lines_per_page（最后一页除外）
        if len(page_lines) < lines_per_page and idx + lines_per_page < len(all_lines):
            page_lines.extend([''] * (lines_per_page - len(page_lines)))
        pages.append(page_lines)
        idx += lines_per_page
    return pages


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def create_source_code_docx(pages_to_submit, total_pages_info):
    """创建格式合规的源代码 Word 文档"""

    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)    # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.orientation = WD_ORIENT.PORTRAIT

    # ---- 页眉 ----
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = HEADER_TEXT
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.style.font.name = HEADER_FONT
    hp.style.font.size = HEADER_FONT_SIZE
    run = hp.runs[0] if hp.runs else hp.add_run(HEADER_TEXT)
    run.font.name = HEADER_FONT
    run.font.size = HEADER_FONT_SIZE
    # 页眉下横线
    pPr = hp._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # ---- 页脚（页码） ----
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 使用 Word 页码域代码
    run1 = fp.add_run('第 ')
    run1.font.name = HEADER_FONT
    run1.font.size = HEADER_FONT_SIZE
    # 页码域
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run2 = fp.add_run()
    run2._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run3 = fp.add_run()
    run3._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run4 = fp.add_run()
    run4._r.append(fldChar2)
    run5 = fp.add_run(' 页')
    run5.font.name = HEADER_FONT
    run5.font.size = HEADER_FONT_SIZE

    # ---- 样式设置 ----
    style = doc.styles['Normal']
    style.font.name = CODE_FONT
    style.font.size = CODE_FONT_SIZE
    style.paragraph_format.line_spacing = 1.0  # 单倍行距
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    # 设置中文字体（兼容）
    style.element.rPr.rFonts.set(qn('w:eastAsia'), HEADER_FONT)

    # ---- 写入代码内容 ----
    # 每页代码用一个段落（每行用换行分隔）
    # 为保证50行/页的精确排版，逐行写入，每行一个段落
    for page_idx, page_lines in enumerate(pages_to_submit):
        for line in page_lines:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_together = True
            # 避免Word自动加段间距
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

            if line == '' or line.strip() == '':
                # 空行：用最小高度的段落
                run = p.add_run(' ')
                run.font.size = CODE_FONT_SIZE
                run.font.name = CODE_FONT
            else:
                run = p.add_run(line)
                run.font.size = CODE_FONT_SIZE
                run.font.name = CODE_FONT
                run.font.color.rgb = None  # 自动颜色（黑）

        # 页间分隔：添加分页符（除了最后一页）
        if page_idx < len(pages_to_submit) - 1:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(docx_break_type='page')  # 分页符

    # ---- 保存 ----
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)

    front_count = min(FRONT_PAGES, total_pages_info['total'])
    back_count = min(BACK_PAGES, total_pages_info['total']) if total_pages_info['total'] > FRONT_PAGES + BACK_PAGES else 0

    print(f'  源代码 Word 文档生成完成！')
    print(f'  文件：{OUTPUT_FILE}')
    print(f'  提交页数：{len(pages_to_submit)} 页')
    print(f'  每页行数：{LINES_PER_PAGE}')
    print(f'  字体：{CODE_FONT} {CODE_FONT_SIZE}')
    print(f'  页边距：2.5cm (四边)')
    print(f'  页眉：{HEADER_TEXT}')
    print(f'  页脚：连续页码')


# 修正：需要导入分页符类型
from docx.oxml.ns import qn as _qn

def add_page_break(doc):
    """添加分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    br = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._r.append(br)


def create_source_code_docx_v2(pages_to_submit, total_pages_info):
    """创建格式合规的源代码 Word 文档（改进版）"""

    from docx.oxml.ns import qn

    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ---- 页眉 ----
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ''
    run = hp.add_run(HEADER_TEXT)
    run.font.name = HEADER_FONT
    run.font.size = HEADER_FONT_SIZE
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 页眉下横线
    pPr = hp._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # ---- 页脚（页码） ----
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ''
    r1 = fp.add_run('第 ')
    r1.font.name = HEADER_FONT
    r1.font.size = HEADER_FONT_SIZE
    # PAGE域代码
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r2 = fp.add_run()
    r2._r.append(fld_begin)
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    r3 = fp.add_run()
    r3._r.append(instr)
    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r4 = fp.add_run()
    r4._r.append(fld_end)
    r5 = fp.add_run(' 页')
    r5.font.name = HEADER_FONT
    r5.font.size = HEADER_FONT_SIZE

    # ---- Normal 样式 ----
    style = doc.styles['Normal']
    style.font.name = CODE_FONT
    style.font.size = CODE_FONT_SIZE
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), CODE_FONT)
    rFonts.set(qn('w:hAnsi'), CODE_FONT)
    rFonts.set(qn('w:eastAsia'), HEADER_FONT)

    # ---- 逐行写入代码 ----
    first_page = True
    for page_idx, page_lines in enumerate(pages_to_submit):
        if not first_page:
            # 分页符
            p_break = doc.add_paragraph()
            p_break.paragraph_format.space_before = Pt(0)
            p_break.paragraph_format.space_after = Pt(0)
            run_break = p_break.add_run()
            br_elem = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
            run_break._r.append(br_elem)
        first_page = False

        for line in page_lines:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            if line.strip() == '':
                run = p.add_run('\u00A0')  # non-breaking space for empty line
                run.font.size = CODE_FONT_SIZE
                run.font.name = CODE_FONT
            else:
                run = p.add_run(line)
                run.font.size = CODE_FONT_SIZE
                run.font.name = CODE_FONT

    # ---- 保存 ----
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)

    print(f'\n  [OK] 源代码 Word 文档生成完成！')
    print(f'  文件：{OUTPUT_FILE}')
    print(f'  提交页数：{len(pages_to_submit)} 页（每页 {LINES_PER_PAGE} 行）')
    print(f'  字体：{CODE_FONT} 小五号(9pt)')
    print(f'  页边距：上下左右各 2.5cm')
    print(f'  页眉："{HEADER_TEXT}"')
    print(f'  页脚：连续页码 "第X页"')
    print(f'  纸张：A4 (21cm × 29.7cm)')


def main():
    print('  步骤1：收集源代码文件...')
    files = collect_source_files()
    print(f'  找到 {len(files)} 个核心源代码文件')

    print('\n  步骤2：读取并清理代码...')
    all_lines, file_markers = build_continuous_code_blocks(files)
    print(f'  有效代码行数：{len(all_lines)}')

    print('\n  步骤3：分页（每页50行）...')
    all_pages = paginate_lines(all_lines, LINES_PER_PAGE)
    total_pages = len(all_pages)
    print(f'  总页数：{total_pages}')

    # 确定提交页范围
    if total_pages <= FRONT_PAGES + BACK_PAGES:
        pages_to_submit = all_pages
        print(f'  总页数 ≤ 60，提交全部 {total_pages} 页')
    else:
        front = all_pages[:FRONT_PAGES]
        back = all_pages[-BACK_PAGES:]
        pages_to_submit = front + back
        print(f'  总页数 > 60，提交前 {FRONT_PAGES} 页 + 后 {BACK_PAGES} 页 = {len(pages_to_submit)} 页')

    total_pages_info = {'total': total_pages}

    print('\n  步骤4：生成 Word 文档...')
    create_source_code_docx_v2(pages_to_submit, total_pages_info)


if __name__ == '__main__':
    main()
