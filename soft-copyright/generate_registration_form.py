# -*- coding: utf-8 -*-
"""
生成"广东省职业技术教育学会2026年职业院校教学智能体创新案例征集活动"
附件2：报名表（申报表）Word文档 - 填写版
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "competition_output")

AGENT_NAME = "码上成长 — Python代码智能评价与学情诊断AI智能体"
SCENE = "代码智能评价与个性化学习辅助"
SUBJECT = "计算机应用/软件技术"
GROUP = "中职组"


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def make_form():
    doc = Document()

    # 页面设置
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    # ===== 大标题 =====
    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t1.add_run("附件2")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("广东省职业技术教育学会2026年职业院校教学智能体\n创新案例征集活动（超星杯）报名表")
    r2.font.size = Pt(16)
    r2.font.bold = True
    r2.font.name = "黑体"
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    # ==================== 主表格 ====================
    # 表格结构：
    # Row 0: 学校 | 联系人 | 手机  (3列)
    # Row 1: 姓名 | 性别 | 出生年月 | 所在部门及职务 | 职称 | 学科 (6列)
    # Row 2: 联系电话 | 电子邮箱 (跨2列)
    # Row 3: 教学智能体情况 (名称 | 应用场景 | 细分学科)
    # Row 4: 活动组别 (中职组)
    # Row 5-9: 团队成员 (5行 × 4列)
    # Row 10: 建设说明
    # Row 11: 知识产权声明
    # Row 12: 成果转化同意
    # Row 13: 盖章

    tbl = doc.add_table(rows=14, cols=6)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    def label(cell, text):
        set_cell_shading(cell, "F0F0F0")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def fill(cell, text, center=False):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def merge_and_fill(row, col_start, col_end, text, center=True):
        row.cells[col_start].merge(row.cells[col_end])
        c = row.cells[col_start]
        fill(c, text, center=center)
        return c

    # ---- Row 0: 学校 / 联系人 / 手机 ----
    label(tbl.rows[0].cells[0], "学校")
    merge_and_fill(tbl.rows[0], 1, 2, "【请填写学校全称】", center=False)
    label(tbl.rows[0].cells[3], "联系人")
    fill(tbl.rows[0].cells[4], "")
    label(tbl.rows[0].cells[5], "手机")
    fill(tbl.rows[0].cells[5], "")

    # ---- Row 1: 个人信息 ----
    label(tbl.rows[1].cells[0], "姓名")
    fill(tbl.rows[1].cells[1], "")
    label(tbl.rows[1].cells[2], "性别")
    fill(tbl.rows[1].cells[3], "")
    label(tbl.rows[1].cells[4], "出生年月")
    fill(tbl.rows[1].cells[5], "")

    # ---- Row 2: 部门/职务/学科 ----
    label(tbl.rows[2].cells[0], "所在部门\n及职务")
    merge_and_fill(tbl.rows[2], 1, 2, "", center=False)
    label(tbl.rows[2].cells[3], "职称")
    fill(tbl.rows[2].cells[4], "")
    label(tbl.rows[2].cells[5], "学科")
    fill(tbl.rows[2].cells[5], "")

    # ---- Row 3: 联系方式 ----
    label(tbl.rows[3].cells[0], "联系电话")
    merge_and_fill(tbl.rows[3], 1, 2, "", center=False)
    label(tbl.rows[3].cells[3], "电子邮箱")
    merge_and_fill(tbl.rows[3], 4, 5, "", center=False)

    # ---- Row 4: 教学智能体情况 ----
    label(tbl.rows[4].cells[0], "教学智\n能体情\n况")
    label(tbl.rows[4].cells[1], "名称")
    merge_and_fill(tbl.rows[4], 2, 3, AGENT_NAME, center=False)
    label(tbl.rows[4].cells[4], "应用场景")
    fill(tbl.rows[4].cells[5], SCENE)
    # 需要加一列给细分学科 - 实际上需要调整
    # 让我重新处理这一行：名称跨2列，然后应用场景，细分学科
    # 先拆分重新来

    # Actually let me just fix row 4 properly by overwriting cells content
    # Row 4 should be: [标签3行] [名称] [名称内容跨2] [应用场景] [细分学科]

    # Let me handle it differently - clear and re-fill row 4
    r4 = tbl.rows[4]
    # Set name
    r4.cells[2].text = ""
    p = r4.cells[2].paragraphs[0]
    p.add_run(AGENT_NAME).font.size = Pt(10)
    # Scene in cell 4
    r4.cells[4].text = ""
    p = r4.cells[4].paragraphs[0]
    p.add_run(SCENE).font.size = Pt(10)
    # Subject in cell 5 - need to add label for it
    r4.cells[5].text = ""
    p = r4.cells[5].paragraphs[0]
    p.add_run(SUBJECT).font.size = Pt(10)

    # Add a small label above subject
    # Actually the original form has: 名称(1col) | 内容(跨2) | 应用场景(1) | 细分学科(1)
    # But our table is 6 columns. Let me adjust.

    # Re-do Row 4 properly: use cells[1] for "名称" label
    r4.cells[1].text = ""
    p = r4.cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("名称")
    run.font.size = Pt(9)
    run.font.bold = True

    # Merge cells[2]+cells[3] for name value
    r4.cells[2].merge(r4.cells[3])
    r4.cells[2].text = ""
    p = r4.cells[2].paragraphs[0]
    run = p.add_run(AGENT_NAME)
    run.font.size = Pt(10)

    # Label "应用场景" in cells[4] 
    r4.cells[4].text = ""
    p = r4.cells[4].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("应用场景\n" + SCENE)
    run.font.size = Pt(9)

    # Label "细分学科" in cells[5] - but we need to fit both label and value
    r4.cells[5].text = ""
    p = r4.cells[5].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("细分学科\n" + SUBJECT)
    run.font.size = Pt(9)

    # ---- Row 5: 活动组别 ----
    label(tbl.rows[5].cells[0], "活动组别")
    group_cell = merge_and_fill(tbl.rows[5], 1, 5,
                                  "☑ 中职组    ☐ 高职组", center=True)

    # ---- Rows 6-10: 团队成员 (5人) ----
    header_row = tbl.rows[6]
    # First column spans rows 6-10 for "其他作者"
    header_row.cells[0].merge(tbl.rows[7].cells[0])
    header_row.cells[0].merge(tbl.rows[8].cells[0])
    header_row.cells[0].merge(tbl.rows[9].cells[0])
    header_row.cells[0].merge(tbl.rows[10].cells[0])
    label(header_row.cells[0], "其他作者\n（个人作者可\n不填团队成\n员）\n\n说明：本表限\n填5人以内")

    # Team member headers
    team_headers = ["排序", "姓名", "职称", "部门及职务", "工作内容"]
    for i, h in enumerate(team_headers):
        c = header_row.cells[i + 1]
        label(c, h)

    # Empty team member rows (rows 7-10 = indices 7,8,9,10)
    for ri in range(7, 11):
        for ci in range(1, 6):
            fill(tbl.rows[ri].cells[ci], "")

    # Add row numbers
    for idx, ri in enumerate(range(7, 11), 1):
        fill(tbl.rows[ri].cells[1], str(idx), center=True)

    # ---- Row 11: 建设说明 ----
    label(tbl.rows[11].cells[0], "建设说明")
    note_cell = merge_and_fill(tbl.rows[11], 1, 5,
                                "（篇幅不够可另附页——详见《附件3_教学智能体建设说明书》）\n\n"
                                "本智能体为\"码上成长 — Python代码智能评价与学情诊断AI智能体\"，"
                                "面向中职Python程序设计课程，集成AI代码评价、在线代码运行、"
                                "AI对话解惑、试题演练、学情统计分析、管理员后台等六大功能模块。"
                                "采用Next.js 16+React 19+DeepSeek大模型技术栈，已在中职实际教学中"
                                "开展一个学期应用实践，取得显著成效。详细技术方案、功能介绍和"
                                "应用成效见附件3说明书。",
                                center=False)

    # ---- Row 12: 知识产权声明 ----
    ip_row = tbl.rows[12]
    label(ip_row.cells[0], "是否保\n证原创")
    ip_content = merge_and_fill(ip_row, 1, 5,
                                 "☑ 是   ☐ 否    \n\n"
                                 "本人保证所申报作品为原创作品，不侵犯任何第三方的知识产权。\n\n"
                                 "签名：________________    日期：______年____月____日",
                                 center=False)

    # ---- Row 13: 成果转化同意 ----
    cv_row = tbl.rows[13]
    label(cv_row.cells[0], "是否同\n意转化")
    cv_content = merge_and_fill(cv_row, 1, 5,
                                 "☑ 是   ☐ 否    \n\n"
                                 "本人同意组委会对作品进行成果转化和推广应用。\n\n"
                                 "签名：________________    日期：______年____月____日",
                                 center=False)

    # ===== 盖章区域（表格外部）=====
    doc.add_paragraph()
    stamp_area = doc.add_paragraph()
    stamp_area.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sr = stamp_area.add_run("学校科技处（盖章）：________________    日期：______年____月____日")
    sr.font.size = Pt(11)
    sr.font.name = "宋体"
    sr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 底部提示
    doc.add_paragraph()
    tip = doc.add_paragraph()
    tr = tip.add_run(
        "填写提示：\n"
        "1. 浅灰色背景单元格为标签栏，白色空白单元格为填写区；\n"
        "2. 「学校」「姓名」等为必填项，请务必完整填写；\n"
        "3. 团队成员按排序 1~5 填写（个人参赛可不填）；\n"
        "4. 填写完成后打印，在「签名」处手写签字，由学校科技处审核盖章；\n"
        "5. 将盖章后的报名表扫描为PDF文件，通过超星平台上传提交。\n"
        "6. 截止日期：2026年8月25日"
    )
    tr.font.size = Pt(9)
    tr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    tr.font.name = "宋体"
    tr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    out_path = os.path.join(OUTPUT_DIR, "附件2_报名表（申报表）.docx")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    fp = make_form()
    print(f"[OK] Registration form generated!")
    print(f"Path: {fp}")
