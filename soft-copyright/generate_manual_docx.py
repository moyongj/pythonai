#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
软著申请 —— 软件说明书 Word 文档生成器
=========================================

根据中国版权保护中心官方规范，生成格式合规的软件用户操作手册 Word 文档：
  - A4 纸，上2.5cm 下2.5cm 左3cm 右2.5cm
  - 页眉：软件全称+版本号
  - 页脚：连续页码 "第X页"
  - 正文：宋体 小四(12pt)，1.5倍行距
  - 章标题：黑体 三号(16pt) 加粗
  - 节标题：黑体 四号(14pt) 加粗
  - 每页 ≥ 30 行（有图除外）
  - 图文并茂，截图占位框标注
  - 功能描述按用户操作路径展开
  - ≥ 15 页

输出：soft-copyright/output/软件说明书.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============ 配置 ============

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, 'output')
OUTPUT_FILE = os.path.join(OUTPUT_DIR, '软件说明书.docx')

SOFTWARE_NAME = '码上成长Python智能学习平台'
VERSION = 'V1.0'
HEADER_TEXT = f'{SOFTWARE_NAME} {VERSION}'
FULL_TITLE = f'{SOFTWARE_NAME} {VERSION} 用户操作手册'

# 字体配置
TITLE_FONT = '黑体'
BODY_FONT = '宋体'
CODE_FONT = 'Courier New'
TITLE_SIZE = Pt(16)      # 三号
SECTION_SIZE = Pt(14)    # 四号
SUBSECTION_SIZE = Pt(12) # 小四
BODY_SIZE = Pt(12)       # 小四
SMALL_SIZE = Pt(10.5)    # 五号
CODE_SIZE = Pt(9)        # 小五号


def setup_page(doc):
    """设置页面、页眉页脚"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ''
    run = hp.add_run(HEADER_TEXT)
    run.font.name = BODY_FONT
    run.font.size = SMALL_SIZE
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = hp._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # 页脚（页码）
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ''
    r1 = fp.add_run('第 ')
    r1.font.name = BODY_FONT
    r1.font.size = SMALL_SIZE
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r2 = fp.add_run()
    r2._r.append(fld_begin)
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    r3 = fp.add_run()
    r3._r.append(instr)
    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r4 = fp.add_run()
    r4._r.append(fld_end)
    r5 = fp.add_run(' 页')
    r5.font.name = BODY_FONT
    r5.font.size = SMALL_SIZE


def add_page_break(doc):
    """添加分页符"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._r.append(br)


def add_title(doc, text, level=0):
    """添加标题"""
    if level == 0:  # 章标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = TITLE_FONT
        run.font.size = TITLE_SIZE
        run.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
    elif level == 1:  # 节标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = TITLE_FONT
        run.font.size = SECTION_SIZE
        run.bold = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
    elif level == 2:  # 小节标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = TITLE_FONT
        run.font.size = SUBSECTION_SIZE
        run.bold = True
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5


def add_body(doc, text, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 2字符缩进
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # 设置中文字体
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), BODY_FONT)
    return p


def add_list_item(doc, text, bullet='●'):
    """添加列表项"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f'{bullet} {text}')
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.74)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), BODY_FONT)


def add_number_item(doc, num, text):
    """添加编号列表项"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f'({num}) {text}')
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.74)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), BODY_FONT)


def add_screenshot_placeholder(doc, description, width_cm=14, height_cm=8):
    """添加截图占位框"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0

    # 用表格模拟占位框
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(width_cm)
    cell.height = Cm(height_cm)

    # 设置边框样式（灰色虚线框表示占位）
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="dashed" w:sz="4" w:space="0" w:color="999999"/>'
        f'  <w:left w:val="dashed" w:sz="4" w:space="0" w:color="999999"/>'
        f'  <w:bottom w:val="dashed" w:sz="4" w:space="0" w:color="999999"/>'
        f'  <w:right w:val="dashed" w:sz="4" w:space="0" w:color="999999"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    # 占位文字
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run(f'[请插入截图：{description}]')
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), BODY_FONT)

    # 删除表格后的空段落
    # (docx 库添加表格后会有空段落)


def add_table_row(table, cells_data, header=False):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = BODY_FONT
        run.font.size = Pt(10.5)
        if header:
            run.bold = True
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), BODY_FONT)


def create_manual():
    """生成完整的软件说明书 Word 文档"""

    doc = Document()
    setup_page(doc)

    # ============ 封面页 ============
    # 空行居中
    for _ in range(6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5

    # 软件名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SOFTWARE_NAME)
    run.font.name = TITLE_FONT
    run.font.size = Pt(22)
    run.bold = True
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), TITLE_FONT)

    # 版本号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(VERSION)
    run.font.name = TITLE_FONT
    run.font.size = Pt(18)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), TITLE_FONT)

    # 空行
    for _ in range(2):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 文档类型
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('用户操作手册')
    run.font.name = TITLE_FONT
    run.font.size = Pt(18)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), TITLE_FONT)

    # 空行
    for _ in range(4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 著作权人信息
    add_body(doc, '著作权人：[请填写著作权人姓名或企业名称]')
    add_body(doc, '开发完成日期：[请填写实际日期]')
    add_body(doc, '首次发表日期：[请填写实际日期或"未发表"]')

    add_page_break(doc)

    # ============ 目录页 ============
    add_title(doc, '目  录', level=0)
    toc_items = [
        ('第一章  软件概述', '1'),
        ('  1.1  开发背景与目的', '1'),
        ('  1.2  适用对象', '1'),
        ('  1.3  功能概述', '2'),
        ('  1.4  技术特点', '2'),
        ('第二章  运行环境', '3'),
        ('  2.1  硬件环境', '3'),
        ('  2.2  软件环境', '3'),
        ('  2.3  网络环境', '3'),
        ('第三章  安装与启动', '4'),
        ('  3.1  安装依赖', '4'),
        ('  3.2  配置环境变量', '4'),
        ('  3.3  启动软件', '5'),
        ('第四章  学生端操作说明', '6'),
        ('  4.1  登录系统', '6'),
        ('  4.2  首页导航', '7'),
        ('  4.3  试题演练', '7'),
        ('  4.4  AI代码评价', '8'),
        ('  4.5  学情统计', '10'),
        ('  4.6  AI解惑对话', '11'),
        ('  4.7  学习资源', '12'),
        ('第五章  管理员端操作说明', '13'),
        ('  5.1  管理员登录', '13'),
        ('  5.2  管理员首页', '13'),
        ('  5.3  学生管理', '14'),
        ('  5.4  题库管理', '15'),
        ('  5.5  学情记录查询', '16'),
        ('  5.6  学情统计分析', '17'),
        ('  5.7  学习资源管理', '18'),
        ('第六章  常见问题与注意事项', '19'),
    ]
    for item_text, page_num in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item_text)
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE
        p.paragraph_format.line_spacing = 1.5

    add_page_break(doc)

    # ============ 第一章 软件概述 ============
    add_title(doc, '第一章  软件概述', level=0)

    add_title(doc, '1.1  开发背景与目的', level=1)
    add_body(doc, '随着人工智能技术的快速发展，编程教育正在成为职业教育的重要组成部分。Python语言因其简洁易学的特点，已成为中职学生编程入门的首选语言。然而，传统的编程教学存在以下问题：教师难以对每位学生的代码进行逐一评价和个性化指导；学生遇到编程困难时缺乏即时帮助；学情数据难以量化追踪和可视化呈现。', indent=True)
    add_body(doc, '"码上成长Python智能学习平台"正是为解决上述问题而开发。软件将AI大模型能力融入编程教学全流程，为学生提供智能代码评价、个性化学习建议和AI对话解惑，为教师提供学情数据可视化分析和题库管理功能，实现编程教学的智能化、个性化和数据化。', indent=True)

    add_title(doc, '1.2  适用对象', level=1)
    add_list_item(doc, '中等职业教育阶段Python编程课程的学生')
    add_list_item(doc, 'Python编程课程任课教师（管理员角色）')
    add_list_item(doc, '职业教育信息化教学研究人员')

    add_title(doc, '1.3  功能概述', level=1)
    add_body(doc, '本软件采用双角色权限设计，分为学生端和管理员端，各自提供不同的功能模块：', indent=True)

    add_body(doc, '学生端功能：', indent=True)
    add_number_item(doc, 1, '试题演练——浏览和练习平台内置的Python编程题目，查看题目要求、提示和示例代码；')
    add_number_item(doc, 2, 'AI代码评价——提交Python代码，AI自动进行多维度智能评价（语法正确性、知识点匹配、代码风格、综合评分）；')
    add_number_item(doc, 3, '学情统计——查看个人历史评价记录、成绩趋势图、知识点掌握分布图，支持生成AI学情分析报告；')
    add_number_item(doc, 4, 'AI解惑对话——与AI智能助手对话，获得Python编程问题的启发式引导和提示；')
    add_number_item(doc, 5, '学习资源——浏览管理员推荐的Python学习资源链接。')

    add_body(doc, '管理员端功能：', indent=True)
    add_number_item(doc, 1, '学生管理——学生信息的增删改查、批量导入、密码重置；')
    add_number_item(doc, 2, '题库管理——Python练习题目的新增、编辑、删除；')
    add_number_item(doc, 3, '学情记录——查看所有学生的代码评价历史记录，支持多条件筛选；')
    add_number_item(doc, 4, '学情分析——班级维度的统计图表（等级分布、分数段分布、知识点频次），支持导出Excel报表；')
    add_number_item(doc, 5, '学习资源管理——管理平台推荐的外链学习资源。')

    add_screenshot_placeholder(doc, '软件功能架构图——展示学生端和管理员端的功能模块划分')

    add_title(doc, '1.4  技术特点', level=1)
    add_number_item(doc, 1, 'AI智能评价——集成OpenAI兼容接口，对学生Python代码进行语法正确性、知识点匹配、代码风格等多维度智能分析，生成结构化评价报告和个性化学习建议；')
    add_number_item(doc, 2, '在线代码运行——内置Python代码执行环境，支持交互式input()输入，代码运行超时保护（10秒），学生可实时验证代码运行效果；')
    add_number_item(doc, 3, '多会话AI对话——支持创建和管理多个对话主题，消息持久化存储，刷新页面后历史消息不丢失，Markdown格式渲染和代码高亮；')
    add_number_item(doc, 4, '数据可视化——使用Recharts图表库展示学情统计数据，包括折线图（成绩趋势）、饼图（等级分布）、柱状图（知识点频次）等；')
    add_number_item(doc, 5, '响应式界面——基于Tailwind CSS设计，适配桌面端和移动端，导航栏和侧边栏支持移动端自适应折叠；')
    add_number_item(doc, 6, '双角色权限——学生和管理员使用同一登录入口，通过权限认证区分角色，登录后自动跳转至对应功能页面；')
    add_number_item(doc, 7, '启发式教学——AI对话采用启发式引导策略，不直接给出完整答案，而是通过提示和引导帮助学生独立解决问题。')

    add_page_break(doc)

    # ============ 第二章 运行环境 ============
    add_title(doc, '第二章  运行环境', level=0)

    add_title(doc, '2.1  硬件环境', level=1)

    # 硬件环境表格
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '项目'
    hdr[1].text = '要求'
    for p in hdr[0].paragraphs:
        for r in p.runs:
            r.bold = True
    for p in hdr[1].paragraphs:
        for r in p.runs:
            r.bold = True
    hw_data = [
        ('CPU', '双核及以上处理器'),
        ('内存', '4GB及以上'),
        ('硬盘', '500MB可用空间（不含数据库增长空间）'),
        ('网络', '宽带互联网连接（用于AI服务调用）'),
    ]
    for item in hw_data:
        add_table_row(table, item)

    add_title(doc, '2.2  软件环境', level=1)
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '项目'
    hdr2[1].text = '要求'
    for p in hdr2[0].paragraphs:
        for r in p.runs:
            r.bold = True
    for p in hdr2[1].paragraphs:
        for r in p.runs:
            r.bold = True
    sw_data = [
        ('操作系统', 'Windows 10及以上 / macOS 10.15及以上 / Linux发行版'),
        ('Node.js', '版本18.17及以上（推荐20.x）'),
        ('pnpm', '版本9.0及以上（包管理器）'),
        ('Python', '3.8及以上（用于在线代码运行功能）'),
        ('浏览器', 'Chrome 90+ / Edge 90+ / Firefox 88+ / Safari 14+'),
        ('数据库', 'SQLite 3.x（软件内置，无需独立安装）'),
    ]
    for item in sw_data:
        add_table_row(table2, item)

    add_title(doc, '2.3  网络环境', level=1)
    add_list_item(doc, '需要能够访问AI服务API（支持OpenAI API及兼容接口如ChatGLM）；')
    add_list_item(doc, '若在内网环境使用，需配置AI服务的公网访问代理或本地化部署模型；')
    add_list_item(doc, '软件运行需要稳定的互联网连接，AI评价和对话功能依赖外部AI服务。')

    add_page_break(doc)

    # ============ 第三章 安装与启动 ============
    add_title(doc, '第三章  安装与启动', level=0)

    add_title(doc, '3.1  安装依赖', level=1)
    add_number_item(doc, 1, '将软件部署包解压至目标目录；')
    add_number_item(doc, 2, '打开系统命令行（终端），进入软件根目录；')
    add_number_item(doc, 3, '执行命令安装项目依赖：pnpm install（注意：必须使用pnpm，不可用npm或yarn）；')
    add_number_item(doc, 4, '等待依赖安装完成（通常需要1~3分钟，取决于网络速度）。')

    add_screenshot_placeholder(doc, '终端执行pnpm install的命令行截图')

    add_title(doc, '3.2  配置环境变量', level=1)
    add_body(doc, '在软件根目录创建或修改".env"文件，按以下说明填写配置项：', indent=True)

    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '配置项'
    hdr3[1].text = '说明'
    hdr3[2].text = '示例'
    for cell in hdr3:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    env_data = [
        ('OPENAI_API_KEY', 'AI服务API密钥', 'sk-xxxxxxxx'),
        ('OPENAI_BASE_URL', 'AI服务接口地址（可选）', 'https://api.openai.com/v1'),
        ('OPENAI_MODEL', '使用的AI模型名称（可选）', 'gpt-4o-mini'),
        ('HOSTNAME', '服务监听地址（可选）', '0.0.0.0'),
        ('PORT', '服务监听端口（可选）', '3000'),
    ]
    for item in env_data:
        add_table_row(table3, item)

    add_title(doc, '3.3  启动软件', level=1)
    add_body(doc, '开发模式启动（推荐用于调试测试）：', indent=True)
    add_number_item(doc, 1, '执行命令：pnpm run dev；')
    add_number_item(doc, 2, '启动后访问：http://localhost:3000；')
    add_number_item(doc, 3, '终端显示"Ready in ..."即表示启动成功。')

    add_body(doc, '生产模式启动（推荐用于正式部署）：', indent=True)
    add_number_item(doc, 1, '构建命令：pnpm run build；')
    add_number_item(doc, 2, '启动命令：pnpm run start；')
    add_number_item(doc, 3, '启动后访问：http://服务器IP:端口。')

    add_screenshot_placeholder(doc, '软件启动后终端输出的成功启动截图')

    add_body(doc, '软件首次启动时自动完成以下初始化工作：', indent=True)
    add_number_item(doc, 1, '在data/目录下创建SQLite数据库文件（mscz.db）；')
    add_number_item(doc, 2, '创建默认管理员账号（用户名：myj，密码：123456）；')
    add_number_item(doc, 3, '创建默认学生账号（学号：01、02、03……，密码：123456）；')
    add_number_item(doc, 4, '导入默认Python练习题库（约60道题目）；')
    add_number_item(doc, 5, '导入默认学习资源链接。')

    add_page_break(doc)

    # ============ 第四章 学生端操作说明 ============
    add_title(doc, '第四章  学生端操作说明', level=0)

    add_title(doc, '4.1  登录系统', level=1)
    add_body(doc, '学生在浏览器地址栏输入软件访问地址后，进入登录页面。登录页面提供"学生登录"和"管理员登录"两种模式，默认显示学生登录界面。', indent=True)
    add_number_item(doc, 1, '在"学号"输入框中输入学号（默认学生账号为01、02、03等）；')
    add_number_item(doc, 2, '在"密码"输入框中输入密码（默认密码为123456）；')
    add_number_item(doc, 3, '点击"登录"按钮提交登录请求；')
    add_number_item(doc, 4, '登录成功后，系统自动跳转至"试题演练"页面；')
    add_number_item(doc, 5, '若账号或密码错误，页面顶部显示红色错误提示"用户名或密码错误"；')
    add_number_item(doc, 6, '若输入为空，提示"请输入用户名和密码"；')
    add_number_item(doc, 7, '若此前已登录且Session未过期，访问登录页会自动跳转至对应功能页面。')

    add_screenshot_placeholder(doc, '学生登录页面截图——展示登录表单（学号输入框、密码输入框、登录按钮）')

    add_title(doc, '4.2  首页导航', level=1)
    add_body(doc, '登录成功后，学生进入"试题演练"页面（学生端默认首页）。页面顶部为导航栏，包含以下元素：', indent=True)
    add_list_item(doc, '左侧：平台Logo图标和名称"码上成长"；')
    add_list_item(doc, '右侧导航菜单按钮：试题演练、代码评价、学情统计、AI解惑、学习资源；')
    add_list_item(doc, '最右侧：当前登录学生姓名显示和"退出登录"按钮。')

    add_body(doc, '学生可通过点击导航栏的各菜单按钮切换至对应功能页面，也可在首页直接点击功能入口卡片快速进入各模块。', indent=True)

    add_screenshot_placeholder(doc, '学生端首页截图——展示导航栏和功能入口卡片')

    add_title(doc, '4.3  试题演练', level=1)
    add_body(doc, '学生点击导航栏"试题演练"按钮后，进入试题演练页面。页面以卡片列表形式展示所有可练习的Python编程题目。', indent=True)

    add_number_item(doc, 1, '页面展示题目卡片列表，每张卡片显示题目标题；')
    add_number_item(doc, 2, '点击任意题目卡片，卡片展开显示题目详情，包括：题目要求（题目描述文字）、提示信息（解题思路提示）、示例代码（参考答案代码）；')
    add_number_item(doc, 3, '学生可参考示例代码，在"代码评价"页面进行实际编写和提交评价；')
    add_number_item(doc, 4, '题目列表支持滚动浏览，题目数据由管理员在后台维护。')

    add_screenshot_placeholder(doc, '试题演练页面截图——展示题目卡片列表和题目详情展开状态')

    add_title(doc, '4.4  AI代码评价（核心功能）', level=1)
    add_body(doc, 'AI代码评价是本软件的核心功能。学生将编写的Python代码提交后，系统调用AI大模型对代码进行全方位智能分析，生成多维度评价报告。', indent=True)

    add_title(doc, '4.4.1  提交代码评价', level=2)
    add_number_item(doc, 1, '点击导航栏"代码评价"按钮进入代码评价页面；')
    add_number_item(doc, 2, '在"选择题目"下拉框中选择要练习的题目（必选，下拉框列出平台所有题目）；')
    add_number_item(doc, 3, '在中央大型代码编辑区输入或粘贴Python代码；')
    add_number_item(doc, 4, '点击"提交评价"按钮；')
    add_number_item(doc, 5, '等待AI分析（通常3~10秒），页面显示加载动画；')
    add_number_item(doc, 6, 'AI评价完成后，页面展示完整的评价报告。')

    add_screenshot_placeholder(doc, '代码评价页面截图——展示题目选择下拉框、代码编辑区、提交按钮')

    add_title(doc, '4.4.2  AI评价报告内容', level=2)
    add_body(doc, 'AI评价报告包含以下内容：', indent=True)
    add_number_item(doc, 1, '评价维度分项展示：题目理解与实现（评分及评价说明）、逻辑思路（评分及评价说明）、代码可读性（评分及评价说明）、语法掌握（评分及评价说明）；')
    add_number_item(doc, 2, '综合评分：百分制总分（0~100），以及对应等级（优秀/良好/及格/待提高）；')
    add_number_item(doc, 3, '知识点识别：AI自动识别代码中涉及的Python知识点，如变量定义、条件语句、循环结构、函数调用等；')
    add_number_item(doc, 4, '改进建议：针对当前代码的个性化学习建议和后续提升方向。')

    add_screenshot_placeholder(doc, 'AI评价报告截图——展示评价维度、综合评分、知识点识别、改进建议')

    add_title(doc, '4.4.3  在线代码运行', level=2)
    add_body(doc, '学生在代码评价页面还可以在线运行Python代码，实时查看运行结果：', indent=True)
    add_number_item(doc, 1, '在代码编辑区编写Python代码后，点击"运行代码"按钮；')
    add_number_item(doc, 2, '若代码包含input()函数需要用户输入，弹出输入框供学生填写输入内容；')
    add_number_item(doc, 3, '系统在后台执行Python代码，将标准输出和错误信息返回至页面；')
    add_number_item(doc, 4, '运行超时时间设为10秒，防止死循环导致服务器阻塞；')
    add_number_item(doc, 5, '运行结果在代码编辑区下方单独的输出面板中展示。')

    add_screenshot_placeholder(doc, '在线代码运行截图——展示代码运行输出结果面板')

    add_page_break(doc)

    add_title(doc, '4.5  学情统计', level=1)
    add_body(doc, '学生点击导航栏"学情统计"按钮后，进入个人学情统计页面。页面展示以下内容：', indent=True)

    add_title(doc, '4.5.1  统计概览卡片', level=2)
    add_number_item(doc, 1, '总提交次数——累计提交代码评价的总次数；')
    add_number_item(doc, 2, '平均分数——所有评价记录的平均得分；')
    add_number_item(doc, 3, '最高分数——历次评价中的最高得分；')
    add_number_item(doc, 4, '掌握知识点数量——AI识别到的不同知识点总数。')

    add_screenshot_placeholder(doc, '学情统计概览卡片截图——展示4个统计指标卡片')

    add_title(doc, '4.5.2  成绩趋势图', level=2)
    add_body(doc, '页面以折线图展示学生的成绩变化趋势：横轴为提交次序（或时间），纵轴为分数（0~100），每条数据点对应一次代码评价记录。通过折线图，学生可以直观地看到自己的学习进步情况或波动趋势。', indent=True)

    add_screenshot_placeholder(doc, '成绩趋势折线图截图——展示分数随时间的变化曲线')

    add_title(doc, '4.5.3  知识点掌握分布', level=2)
    add_body(doc, '页面以柱状图展示各Python知识点的掌握情况，颜色区分不同掌握程度：掌握（绿色）、部分掌握（黄色）、未掌握（红色）。学生可以据此识别自己的薄弱知识点，重点加强练习。', indent=True)

    add_screenshot_placeholder(doc, '知识点掌握分布柱状图截图')

    add_title(doc, '4.5.4  历史评价记录', level=2)
    add_body(doc, '页面下方展示历史评价记录列表，每条记录显示：题目名称、提交时间、分数、等级。点击任意记录可展开查看完整的AI评价报告详情。', indent=True)

    add_screenshot_placeholder(doc, '历史评价记录列表截图')

    add_title(doc, '4.5.5  生成AI学情分析报告', level=2)
    add_body(doc, '学生可点击"生成分析报告"按钮，AI将基于历史评价记录自动生成个性化学习分析报告，内容包括：学习总结（整体表现评价）、优势亮点（做得好的方面）、薄弱环节（需要加强的知识点）、改进建议（具体的学习建议）、后续学习规划（阶段性学习目标）。报告以美观的HTML页面展示，支持中文排版。', indent=True)

    add_screenshot_placeholder(doc, 'AI学情分析报告截图——展示生成的个性化分析报告内容')

    add_page_break(doc)

    add_title(doc, '4.6  AI解惑对话', level=1)
    add_body(doc, '学生点击导航栏"AI解惑"按钮后，进入AI对话页面。该功能支持学生与AI智能助手进行对话式交流，询问Python编程问题并获得启发式引导。', indent=True)

    add_title(doc, '4.6.1  会话管理', level=2)
    add_number_item(doc, 1, '页面左侧为"会话列表"面板，展示所有已创建的对话主题；')
    add_number_item(doc, 2, '点击"新建对话"按钮创建新的对话主题，系统自动命名（如"新对话1"）；')
    add_number_item(doc, 3, '点击已有会话可切换至该对话，查看历史消息；')
    add_number_item(doc, 4, '支持重命名会话标题、删除不再需要的会话；')
    add_number_item(doc, 5, '所有对话消息持久化存储在数据库中，刷新页面或重新登录后历史消息不丢失。')

    add_screenshot_placeholder(doc, 'AI对话页面截图——展示左侧会话列表和右侧对话区域')

    add_title(doc, '4.6.2  对话交互', level=2)
    add_number_item(doc, 1, '在页面底部输入框中输入Python编程问题；')
    add_number_item(doc, 2, '点击"发送"按钮或按Ctrl+Enter发送消息；')
    add_number_item(doc, 3, 'AI回复实时显示在对话区域，支持Markdown格式渲染（代码高亮、标题、列表等）；')
    add_number_item(doc, 4, 'AI助手严格限制只回答Python相关的编程问题，对于非Python问题会礼貌拒绝并引导；')
    add_number_item(doc, 5, 'AI采用启发式教学策略，不直接给出完整答案，而是通过提示和引导帮助学生独立解决问题。')

    add_screenshot_placeholder(doc, 'AI对话交互截图——展示学生提问和AI回复的对话内容')

    add_title(doc, '4.6.3  移动端适配', level=2)
    add_body(doc, '在移动端设备上，会话列表隐藏为可折叠的侧边栏。学生点击左上角的菜单按钮可展开或收起会话列表，对话区域自动占满屏幕宽度，操作体验与桌面端一致。', indent=True)

    add_screenshot_placeholder(doc, 'AI对话移动端截图——展示折叠侧边栏和对话区域')

    add_title(doc, '4.7  学习资源', level=1)
    add_body(doc, '学生点击导航栏"学习资源"按钮后，进入学习资源页面。页面以卡片形式展示管理员推荐的Python学习资源链接。', indent=True)
    add_number_item(doc, 1, '每张资源卡片显示：资源标题、资源描述、资源链接、资源图标；')
    add_number_item(doc, 2, '点击"访问资源"按钮，在内嵌窗口中直接打开资源链接；')
    add_number_item(doc, 3, '点击"外部打开"按钮，在浏览器新标签页中打开资源链接；')
    add_number_item(doc, 4, '学习资源由管理员在后台维护，学生端只可浏览和访问。')

    add_screenshot_placeholder(doc, '学习资源页面截图——展示资源卡片列表和资源链接访问')

    add_page_break(doc)

    # ============ 第五章 管理员端操作说明 ============
    add_title(doc, '第五章  管理员端操作说明', level=0)

    add_title(doc, '5.1  管理员登录', level=1)
    add_body(doc, '管理员在登录页面点击"管理员登录"按钮切换至管理员登录界面，输入管理员账号和密码后点击"登录"按钮。登录成功后系统自动跳转至管理员首页（/admin）。', indent=True)

    add_screenshot_placeholder(doc, '管理员登录界面截图——展示管理员登录表单')

    add_title(doc, '5.2  管理员首页', level=1)
    add_body(doc, '管理员首页展示平台运行概况，包括：学生总数统计、题目总数统计、评价记录总数统计、今日新增评价次数。页面还提供快捷操作入口，管理员可点击相应按钮快速跳转至各管理模块。', indent=True)

    add_screenshot_placeholder(doc, '管理员首页截图——展示统计概览和快捷操作入口')

    add_title(doc, '5.3  学生管理', level=1)
    add_body(doc, '管理员点击左侧菜单"学生管理"后，进入学生管理页面。页面展示学生列表表格，支持以下操作：', indent=True)

    add_number_item(doc, 1, '新增学生：点击"新增学生"按钮，在弹出表单中填写学号和姓名，系统自动生成默认密码（123456），点击"确认"完成创建；')
    add_number_item(doc, 2, '编辑学生：点击某行学生记录的"编辑"按钮，修改姓名后点击"保存"；')
    add_number_item(doc, 3, '删除学生：点击某行学生记录的"删除"按钮，在确认弹窗中点击"确认"删除（注意：删除学生将同时删除该学生的所有评价记录）；')
    add_number_item(doc, 4, '重置密码：点击某行学生记录的"重置密码"按钮，密码将被重置为默认密码（123456）；')
    add_number_item(doc, 5, '批量导入：点击"批量导入"按钮，下载Excel模板，填写学生信息后上传，系统自动批量创建学生账号。')

    add_screenshot_placeholder(doc, '学生管理页面截图——展示学生列表表格和操作按钮')
    add_screenshot_placeholder(doc, '新增学生弹出表单截图——展示学号和姓名输入框')

    add_title(doc, '5.4  题库管理', level=1)
    add_body(doc, '管理员点击左侧菜单"题库管理"后，进入题库管理页面。页面展示题目列表表格，支持以下操作：', indent=True)

    add_number_item(doc, 1, '新增题目：点击"新增题目"按钮，在弹出表单中填写题目标题、题目要求、提示信息、示例代码，点击"确认"完成创建；')
    add_number_item(doc, 2, '编辑题目：点击某行题目记录的"编辑"按钮，修改题目信息后点击"保存"；')
    add_number_item(doc, 3, '删除题目：点击某行题目记录的"删除"按钮，确认后删除；')
    add_number_item(doc, 4, '查看题目：列表中每行显示题号、题目标题、创建时间和操作按钮。')

    add_screenshot_placeholder(doc, '题库管理页面截图——展示题目列表和新增/编辑表单')

    add_title(doc, '5.5  学情记录查询', level=1)
    add_body(doc, '管理员点击左侧菜单"学情记录"后，进入学情记录查询页面。页面展示所有学生的代码评价历史记录列表。', indent=True)

    add_number_item(doc, 1, '列表每行显示：学生姓名、题目名称、综合评分、等级、提交时间、操作按钮；')
    add_number_item(doc, 2, '筛选功能：按学生姓名筛选（下拉框）、按题目筛选（下拉框）、按等级筛选（下拉框）；')
    add_number_item(doc, 3, '点击"查询"按钮执行筛选，点击"重置"按钮清除筛选条件；')
    add_number_item(doc, 4, '点击某条记录的"查看详情"按钮，弹出详情面板展示完整的AI评价报告。')

    add_screenshot_placeholder(doc, '学情记录查询页面截图——展示记录列表和筛选条件')
    add_screenshot_placeholder(doc, '评价记录详情面板截图——展示完整的AI评价报告')

    add_page_break(doc)

    add_title(doc, '5.6  学情统计分析', level=1)
    add_body(doc, '管理员点击左侧菜单"学情分析"后，进入学情统计分析页面。页面以多种图表展示班级维度的学情数据：', indent=True)

    add_number_item(doc, 1, '平均分统计卡片——展示全班平均分数；')
    add_number_item(doc, 2, '等级分布饼图——展示"优秀/良好/及格/待提高"各等级的人数占比；')
    add_number_item(doc, 3, '分数段分布柱状图——展示各分数段（0~59/60~69/70~79/80~89/90~100）的学生人数；')
    add_number_item(doc, 4, '知识点频次条形图——展示各Python知识点在评价记录中出现的频次排序，帮助教师了解学生普遍薄弱的知识点；')
    add_number_item(doc, 5, '导出功能：点击"导出Excel"按钮，系统生成包含所有学情记录的Excel文件并自动下载。')

    add_screenshot_placeholder(doc, '学情统计分析页面截图——展示等级分布饼图和分数段柱状图')
    add_screenshot_placeholder(doc, '知识点频次条形图截图——展示各知识点出现频次排序')

    add_title(doc, '5.7  学习资源管理', level=1)
    add_body(doc, '管理员点击左侧菜单"学习资源"后，进入学习资源管理页面。页面展示资源列表表格，支持以下操作：', indent=True)

    add_number_item(doc, 1, '新增资源：点击"新增资源"按钮，填写资源名称、描述、链接地址、排序序号，点击"确认"完成创建；')
    add_number_item(doc, 2, '编辑资源：点击某行资源的"编辑"按钮，修改信息后点击"保存"；')
    add_number_item(doc, 3, '删除资源：点击某行资源的"删除"按钮，确认后删除；')
    add_number_item(doc, 4, '排序调整：修改排序序号后，列表自动按序号升序排列，序号越小的资源在学生端显示越靠前。')

    add_screenshot_placeholder(doc, '学习资源管理页面截图——展示资源列表和编辑表单')

    add_page_break(doc)

    # ============ 第六章 常见问题与注意事项 ============
    add_title(doc, '第六章  常见问题与注意事项', level=0)

    add_title(doc, '6.1  常见问题', level=1)

    add_title(doc, 'Q1：学生忘记密码怎么办？', level=2)
    add_body(doc, '由管理员在"学生管理"模块中对该学生执行"重置密码"操作，密码将重置为默认密码（123456）。', indent=True)

    add_title(doc, 'Q2：AI评价一直没有响应怎么办？', level=2)
    add_list_item(doc, '确认.env文件中OPENAI_API_KEY配置正确；')
    add_list_item(doc, '确认服务器能够访问AI服务接口地址（网络连通性）；')
    add_list_item(doc, '检查AI服务账户是否有足够余额；')
    add_list_item(doc, '若AI服务响应超时，可尝试更换模型或切换AI服务供应商。')

    add_title(doc, 'Q3：在线运行Python代码时报错怎么办？', level=2)
    add_body(doc, '请在服务器上安装Python（3.8及以上版本），并确保python或python3命令可在命令行中直接执行。软件会自动检测Python路径。', indent=True)

    add_title(doc, 'Q4：如何备份数据？', level=2)
    add_body(doc, '平台的全部数据存储在data/mscz.db文件中，定期复制该文件至安全位置即可完成数据备份。建议每周备份一次。', indent=True)

    add_title(doc, 'Q5：如何修改管理员密码？', level=2)
    add_body(doc, '管理员可通过数据库工具直接修改管理员密码字段，或联系开发人员协助修改。生产环境部署时务必修改默认密码。', indent=True)

    add_title(doc, '6.2  注意事项', level=1)
    add_number_item(doc, 1, '本软件为教育辅助工具，AI评价结果仅供参考，不作为正式考试成绩依据；')
    add_number_item(doc, 2, '使用AI功能需要有效的AI服务API密钥，请妥善保管，不要泄露给他人；')
    add_number_item(doc, 3, '建议定期备份数据库文件（data/mscz.db），防止数据丢失；')
    add_number_item(doc, 4, '若部署在公网服务器，请务必修改默认管理员密码和默认学生密码，确保系统安全；')
    add_number_item(doc, 5, '本软件使用的AI服务由第三方提供，服务可用性和评价质量取决于AI服务供应商；')
    add_number_item(doc, 6, '软件内置题库可根据教学需要随时由管理员增删改查，建议定期更新题目以保持教学内容与课程进度同步；')
    add_number_item(doc, 7, '退出登录时请点击导航栏右侧的"退出登录"按钮，不要直接关闭浏览器窗口，以确保Session正确失效。')

    # ---- 保存 ----
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)

    print(f'\n  [OK] 软件说明书 Word 文档生成完成！')
    print(f'  文件：{OUTPUT_FILE}')
    print(f'  格式：A4纸，宋体小四正文，黑体标题，1.5倍行距')
    print(f'  页眉："{HEADER_TEXT}"')
    print(f'  页脚：连续页码 "第X页"')
    print(f'  注意：文档中包含截图占位框，请替换为实际软件截图后再提交！')


if __name__ == '__main__':
    create_manual()
